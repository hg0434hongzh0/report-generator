import tkinter as tk
from tkinter import ttk, messagebox
from tkinterdnd2 import TkinterDnD
from docx import Document
from docx.shared import Inches
from PIL import Image, ImageTk, ImageGrab
import os

# 修复建议字典
fix_suggestions = {
    "跨站脚本攻击(XSS)": "对用户输入进行严格的验证和编码，避免注入脚本。使用安全的内容安全策略(CSP)来防止不受信任的代码执行，并启用HTTPOnly和Secure标记来保护cookie。",
    "配置错误": "检查并修复错误的系统配置，确保按照最佳安全实践进行配置。定期进行安全评估，使用自动化工具检测配置错误，并启用最小权限原则。",
    "弱口令": "要求用户使用强密码策略，定期更改密码。建议采用双因素认证(2FA)，并设置账户锁定机制防止暴力破解。",
    "疑似被黑": "立即采取行动，分析入侵点，并修复所有被攻击的漏洞。建议检查系统日志、更新所有软件版本，并更换所有系统密码。",
    "任意文件上传(GetShell)": "限制上传的文件类型和大小，检查并处理上传的文件内容。启用MIME类型验证和文件扫描，确保文件不包含恶意代码。",
    "信息泄露": "加密敏感数据，严格控制数据的访问权限。使用加密传输协议（如TLS）保护网络通信，并定期清理旧的或不必要的数据。",
    "存在后门": "删除后门程序，并检查系统中是否存在其他恶意软件。建议对系统进行全面病毒扫描，重置所有访问凭证，并启用入侵检测系统(IDS)。",
    "逻辑漏洞": "检查业务逻辑，确保符合设计要求并排除任何安全漏洞。建议对所有输入进行验证，使用自动化测试工具进行漏洞扫描，并进行代码审查。",
    "RCE": "确保代码执行安全，限制不信任输入的执行权限。建议禁用不必要的代码执行功能，限制执行环境，使用沙箱隔离高风险操作。",
    "SQL注入": "使用参数化查询或ORM，避免直接拼接SQL语句。定期检查和修复数据库权限，启用数据库的输入过滤功能。",
    "解析漏洞": "确保输入的文件或数据格式符合预期，并进行严格的格式验证。禁用未使用的解析器，采用白名单策略来验证文件格式。",
    "URL重定向": "对用户输入的URL进行验证和限制，避免重定向到恶意地址。建议使用硬编码安全URL，并对URL参数进行严格校验。",
    "权限失控(越权)": "检查权限控制逻辑，确保资源只能被授权用户访问。建议引入基于角色的访问控制(RBAC)或基于属性的访问控制(ABAC)，并定期审查权限配置。",
    "验证码绕过": "增加验证码复杂性，并进行行为分析防止自动化攻击。建议结合用户行为模式分析，限制短时间内的重复尝试，启用图形或多步验证的验证码。",
    "服务器端请求伪造(SSRF)": "限制服务器发送的请求，避免请求不必要的内部资源。启用网络隔离机制，限制服务器的外部网络访问，使用黑名单或白名单策略。",
    "跨站请求伪造(CSRF)": "使用CSRF令牌，确保请求的合法性。建议引入双重身份验证，并确保敏感操作需要重新输入凭证。",
    "路径遍历": "验证和规范化文件路径，避免访问不应公开的文件。建议禁用相对路径访问，并将用户文件操作限定在指定目录中。",
    "任意文件下载": "对下载的文件路径进行验证，确保只能下载预期范围内的文件。建议对下载操作进行日志记录，并启用下载限速机制。",
    "任意文件读取": "检查并限制文件读取权限，确保只能访问允许的文件。建议禁用文件读取API的暴露，并对重要文件进行权限隔离。",
    "任意文件删除": "严格限制文件删除操作，确保只有授权用户可以删除文件。建议启用回收站或文件恢复机制，并记录所有删除操作。",
    "XML外部实体注入(XXE)": "禁用XML解析中的外部实体，避免外部实体注入攻击。建议使用安全的XML解析库，并对外部实体进行严格的访问控制。",
    "失效的身份认证": "强化身份认证机制，采用双因素认证等措施。建议定期审查认证机制，并通过安全审计确保身份验证流程的有效性。",
    "文件包含": "对文件路径进行验证，避免加载不安全的文件。建议使用白名单机制，只允许指定的文件路径，禁用动态文件加载功能。",
    "条件竞争": "加锁关键资源，确保多线程访问安全。建议使用事务锁或同步机制，并对代码进行并发性测试。",
    "可预测漏洞": "增加随机性，避免使用可预测的值。建议引入强加密的随机数生成器，并定期轮换使用的密钥或凭证。",
    "硬编码": "避免在代码中硬编码敏感信息，使用安全的凭证管理方案。建议使用环境变量或加密密钥库，并确保代码中不留敏感数据。",
    "域传送": "限制DNS服务器的域传送功能，避免敏感信息泄露。建议启用区域传输控制，只允许可信IP执行域传送。",
    "拒绝服务": "增加防护措施，限制每个IP的请求速率，避免DDoS攻击。建议引入Web应用防火墙(WAF)，并启用自动化的流量监控和响应策略。",
    "其他": "根据具体情况采取相应的安全措施。建议结合现有的行业最佳实践，进行全方位的安全评估和漏洞修复。"
}


# 全局变量，用于存储图片
pasted_images = []


# 处理粘贴事件
def handle_paste(event=None):
    global pasted_images

    try:
        # 从剪贴板获取图像
        image = ImageGrab.grabclipboard()

        if isinstance(image, Image.Image):
            # 在Tkinter文本框中显示图片
            tk_image = ImageTk.PhotoImage(image)
            image_label = tk.Label(reproduction_text, image=tk_image)
            image_label.image = tk_image  # 防止被垃圾回收
            reproduction_text.window_create(tk.END, window=image_label)

            # 存储图片到列表，稍后生成报告时使用
            pasted_images.append(image)
        else:
            print("剪贴板中没有图像")
    except Exception as e:
        print(f"处理粘贴时出错: {e}")


# 自动填充修复建议
def update_solution(event):
    selected_vuln = category_combobox.get()
    solution_text.delete("1.0", tk.END)  # 清空修复建议文本框
    solution_text.insert(tk.END, fix_suggestions.get(selected_vuln, "根据具体情况采取相应的安全措施。"))


def submit_report():
    vendor = vendor_entry.get()
    src = src_entry.get()
    title = title_entry.get()
    category = category_combobox.get()
    size= size_combobox.get()
    level = level_combobox.get()
    url = url_entry.get()
    website_name = website_name_entry.get()
    website_ip = website_ip_entry.get()
    industry = industry_combobox.get()
    region = region_combobox.get()
    description = description_text.get("1.0", tk.END).strip()
    reproduction = reproduction_text.get("1.0", tk.END).strip()
    solution = solution_text.get("1.0", tk.END).strip()

    if not (
            vendor and src and title and category and level and url and website_name and website_ip and industry and region and description and reproduction and solution):
        messagebox.showerror("错误", "请填写所有必填项！")
        return

    # 创建docx文档
    doc = Document()
    doc.add_heading('漏洞报告', level=1)

    doc.add_paragraph(f"厂商名称: {vendor}")
    doc.add_paragraph(f"项目名称: {src}")
    doc.add_paragraph(f"漏洞标题: {title}")
    doc.add_paragraph(f"漏洞类型: {category}")
    doc.add_paragraph(f"漏洞类别: {size}")
    doc.add_paragraph(f"漏洞等级: {level}")
    doc.add_paragraph(f"漏洞URL: {url}")
    doc.add_paragraph(f"网站名称: {website_name}")
    doc.add_paragraph(f"网站IP: {website_ip}")
    doc.add_paragraph(f"所属行业: {industry}")
    doc.add_paragraph(f"所属地区: {region}")
    doc.add_paragraph(f"漏洞描述:\n{description}")
    doc.add_paragraph(f"复现步骤:\n{reproduction}")
    doc.add_paragraph(f"修复方案:\n{solution}")

    # 将粘贴的图片插入到报告中
    for idx, img in enumerate(pasted_images):
        img_path = f'image_{idx}.png'
        img.save(img_path)  # 保存图片
        doc.add_paragraph(f"复现步骤截图 {idx + 1}:")
        doc.add_picture(img_path, width=Inches(4.0))  # 插入图片

        # 可选：在插入图片后删除本地图片文件
        os.remove(img_path)

    # 保存文档
    doc.save("报告.docx")

    messagebox.showinfo("成功", "报告已生成到 '报告.docx'.")



# 创建主窗口
root = TkinterDnD.Tk()  # 修改为DnD窗口
root.title("代码是开源的，你的心不是(qq:2861395370)")
root.geometry("600x600")

# 创建Canvas并添加滚动条
canvas = tk.Canvas(root)
scrollbar = tk.Scrollbar(root, orient="vertical", command=canvas.yview)
scrollable_frame = tk.Frame(canvas)

scrollable_frame.bind(
    "<Configure>",
    lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
)

canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
canvas.configure(yscrollcommand=scrollbar.set)

# 标题
tk.Label(scrollable_frame, text="HG0434报告生成器", font=('Arial', 16)).grid(row=0, column=0, columnspan=2, pady=10)

# 厂商名称
tk.Label(scrollable_frame, text="厂商名称:", font=('Arial', 12)).grid(row=1, column=0, sticky="w", padx=10, pady=5)
vendor_entry = tk.Entry(scrollable_frame, width=50)
vendor_entry.grid(row=1, column=1, padx=10, pady=5)

# 所属SRC
tk.Label(scrollable_frame, text="项目名称:", font=('Arial', 12)).grid(row=2, column=0, sticky="w", padx=10, pady=5)
src_entry = tk.Entry(scrollable_frame, width=50)
src_entry.grid(row=2, column=1, padx=10, pady=5)

# 漏洞标题
tk.Label(scrollable_frame, text="漏洞标题:", font=('Arial', 12)).grid(row=3, column=0, sticky="w", padx=10, pady=5)
title_entry = tk.Entry(scrollable_frame, width=50)
title_entry.grid(row=3, column=1, padx=10, pady=5)

# 漏洞类别
tk.Label(scrollable_frame, text="漏洞类型:", font=('Arial', 12)).grid(row=4, column=0, sticky="w", padx=10, pady=5)
category_combobox = ttk.Combobox(scrollable_frame, values=list(fix_suggestions.keys()), width=47)
category_combobox.grid(row=4, column=1, padx=10, pady=5)
category_combobox.bind("<<ComboboxSelected>>", update_solution)

# 漏洞类别
tk.Label(scrollable_frame, text="漏洞类别:", font=('Arial', 12)).grid(row=5, column=0, sticky="w", padx=10, pady=5)
size_combobox = ttk.Combobox(scrollable_frame, values=['事件型','通用型'], width=47)
size_combobox.grid(row=5, column=1, padx=10, pady=5)

# 漏洞等级
tk.Label(scrollable_frame, text="漏洞等级:", font=('Arial', 12)).grid(row=6, column=0, sticky="w", padx=10, pady=5)
level_combobox = ttk.Combobox(scrollable_frame, values=['低危', '中危', '高危'], width=47)
level_combobox.grid(row=6, column=1, padx=10, pady=5)

# 漏洞URL
tk.Label(scrollable_frame, text="漏洞URL:", font=('Arial', 12)).grid(row=7, column=0, sticky="w", padx=10, pady=5)
url_entry = tk.Entry(scrollable_frame, width=50)
url_entry.grid(row=7, column=1, padx=10, pady=5)

# 网站名称
tk.Label(scrollable_frame, text="网站名称:", font=('Arial', 12)).grid(row=8, column=0, sticky="w", padx=10, pady=5)
website_name_entry = tk.Entry(scrollable_frame, width=50)
website_name_entry.grid(row=8, column=1, padx=10, pady=5)

# 网站IP
tk.Label(scrollable_frame, text="网站IP:", font=('Arial', 12)).grid(row=9, column=0, sticky="w", padx=10, pady=5)
website_ip_entry = tk.Entry(scrollable_frame, width=50)
website_ip_entry.grid(row=9, column=1, padx=10, pady=5)

# 所属行业
tk.Label(scrollable_frame, text="所属行业:", font=('Arial', 12)).grid(row=10, column=0, sticky="w", padx=10, pady=5)
industry_combobox = ttk.Combobox(scrollable_frame, values=[
    "国家机构", "国际组织", "金融业", "电信运营", "广播电视/卫星传输", "文化传媒",
    "教育", "卫生医药", "公共事业/民生", "交通运输/仓储/邮政(物流)业", "国防科工", "水利",
    "能源/矿业", "化工", "环境保护", "制造业", "农/林/牧/渔业", "互联网/软件和信息技术",
    "建筑/房地产", "消费/住宿/餐饮", "租赁/商务服务", "体育/娱乐/旅游", "社会组织", "军队",
    "车联网", "其他"
], width=47)
industry_combobox.grid(row=10, column=1, padx=10, pady=5)

# 所属地区
tk.Label(scrollable_frame, text="所属地区:", font=('Arial', 12)).grid(row=11, column=0, sticky="w", padx=10, pady=5)
region_combobox = ttk.Combobox(scrollable_frame, values=[
    "北京市", "天津市", "河北省", "山西省", "内蒙古自治区", "辽宁省",
    "吉林省", "黑龙江省", "上海市", "江苏省", "浙江省", "安徽省",
    "福建省", "江西省", "山东省", "河南省", "湖北省", "湖南省",
    "广东省", "广西壮族自治区", "海南省", "重庆市", "四川省", "贵州省",
    "云南省", "西藏自治区", "陕西省", "甘肃省", "青海省", "宁夏回族自治区",
    "新疆维吾尔自治区", "台湾省", "香港特别行政区", "澳门特别行政区"
], width=47)
region_combobox.grid(row=11, column=1, padx=10, pady=5)

# 漏洞描述
tk.Label(scrollable_frame, text="漏洞描述:", font=('Arial', 12)).grid(row=12, column=0, sticky="w", padx=10, pady=5)
description_text = tk.Text(scrollable_frame, width=50, height=5)
description_text.grid(row=12, column=1, padx=10, pady=5)

# 复现步骤
tk.Label(scrollable_frame, text="复现步骤:", font=('Arial', 12)).grid(row=13, column=0, sticky="w", padx=10, pady=5)

# 创建一个框架来容纳文本框和滚动条
text_frame = tk.Frame(scrollable_frame)
text_frame.grid(row=13, column=1, padx=10, pady=5, sticky="nsew")

# 添加滚动条
scrollbar = tk.Scrollbar(text_frame)
scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

# 调大复现步骤文本框，并将其与滚动条绑定
reproduction_text = tk.Text(text_frame, width=70, height=15, yscrollcommand=scrollbar.set)
reproduction_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

# 绑定滚动条与文本框
scrollbar.config(command=reproduction_text.yview)

# 绑定粘贴事件
reproduction_text.bind('<Control-v>', handle_paste)

# 配置列和行的权重，以确保框架在调整大小时也能自适应
scrollable_frame.grid_rowconfigure(13, weight=1)
scrollable_frame.grid_columnconfigure(1, weight=1)


# 修复建议
tk.Label(scrollable_frame, text="修复建议:", font=('Arial', 12)).grid(row=14, column=0, sticky="w", padx=10, pady=5)
solution_text = tk.Text(scrollable_frame, width=50, height=5)
solution_text.grid(row=14, column=1, padx=10, pady=5)

# 提交按钮
submit_button = tk.Button(scrollable_frame, text="生成报告", command=submit_report)
submit_button.grid(row=15, column=0, columnspan=2, pady=20)

# 添加Canvas和Scrollbar
canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

#弹窗-不喜欢可以删掉QWQ
root = tk.Tk()
root.withdraw()  # 隐藏主窗口
messagebox.showinfo("提示", "欢迎使用HG0434报告生成器\n本工具“仿照”360众包报告标准所制作,旨在提升日常做HVV/SRC等(虽然基本都有模板)的规范性\n生成报告及时保存，别被覆盖")
root.destroy()

# 启动主循环
root.mainloop()